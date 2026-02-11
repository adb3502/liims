"""
Refactored label generator for BHARAT Study
Can be called programmatically from Streamlit app
"""
from pathlib import Path
import re
from typing import List, Dict
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Label groups (22 labels per participant, CS2 is empty/blank)
LABEL_GROUPS = {
    'cryovial': ['P1', 'P2', 'P3', 'P4', 'P5'],
    'epigenetics': ['E1', 'E2', 'E3', 'E4'],
    'samples': ['CS1', 'R1', 'H1', ''],  # 4th position is blank for A, H2 for B participants
    'edta': ['EDTA1', 'EDTA2', 'EDTA3', 'EDTA4'],
    'sst_fl_blood': ['SST1', 'SST2', 'Fl1', 'B1']
}

# Cryovial layout (5 per row with row gutters)
CRYO_CONFIG = {
    'rows_per_page': 17,
    'row_h_cm': 1.3,
    'row_gutter_h_cm': 0.3,
    'col_widths_cm': [3.3, 0.3, 3.3, 0.3, 3.3, 0.3, 3.3, 0.3, 3.3],
    'label_col_idx': [0, 2, 4, 6, 8],
    'margins_cm': {'top': 0.5, 'bottom': 0.5, 'left': 0.5, 'right': 0.5},
    'font': 'Aptos',
    'font_size': 11,
    'font_size_header': 7,
    'has_row_gutters': True
}

# Normal layout (4 per row, no row gutters)
NORMAL_CONFIG = {
    'rows_per_page': 21,
    'row_h_cm': 1.25,
    'row_gutter_h_cm': 0,
    'col_widths_cm': [4.6, 0.5, 4.6, 0.5, 4.6, 0.5, 4.6],
    'label_col_idx': [0, 2, 4, 6],
    'margins_cm': {'top': 1.8, 'bottom': 0.5, 'left': 0.5, 'right': 0.5},
    'font': 'Aptos',
    'font_size': 11,
    'font_size_header': 7,
    'has_row_gutters': False
}

A4_W_CM = 21.0
A4_H_CM = 29.7


def cm_to_twips(cm):
    return int(cm * 567)


def is_b_participant(participant: str) -> bool:
    """Check if participant code has 'B' (e.g., 1B-123, 2B-456)."""
    return bool(re.match(r'^\d+B-', participant))


def create_label_collections(participants: List[str]) -> Dict[str, List[str]]:
    """Create collections of labels grouped by document type."""
    collections = {
        'cryovial': [],
        'epigenetics': [],
        'samples': [],
        'edta': [],
        'sst_fl_blood': []
    }

    for participant in participants:
        for group_name, suffixes in LABEL_GROUPS.items():
            for suffix in suffixes:
                if suffix:
                    label_code = f"{participant}-{suffix}"
                    collections[group_name].append(label_code)
                else:
                    # For B participants, add H2 in the empty column; for others, leave blank
                    if is_b_participant(participant):
                        collections[group_name].append(f"{participant}-H2")
                    else:
                        collections[group_name].append('')

    return collections


def set_table_xml_enforced(table, col_widths_cm):
    """Add fixed layout + zero spacing + explicit widths."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Fixed layout
    try:
        layout_xml = parse_xml(r'<w:tblLayout %s w:type="fixed"/>' % nsdecls("w"))
        tblPr.append(layout_xml)
    except Exception:
        pass

    # Zero cell spacing
    try:
        spacing_xml = parse_xml(
            r'<w:tblCellSpacing %s w:w="0" w:type="dxa"/>' % nsdecls("w")
        )
        tblPr.append(spacing_xml)
    except Exception:
        pass

    # Table width
    total_twips = sum(cm_to_twips(c) for c in col_widths_cm)
    try:
        tblW_xml = parse_xml(r'<w:tblW %s w:w="%d" w:type="dxa"/>' % (nsdecls("w"), total_twips))
        tblPr.append(tblW_xml)
    except Exception:
        pass

    # Build tblGrid
    try:
        tblGrid = OxmlElement('w:tblGrid')
        for w in col_widths_cm:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(cm_to_twips(w)))
            tblGrid.append(gridCol)
        tbl.append(tblGrid)
    except Exception:
        pass


def set_cell_tcPr_width(cell, width_cm):
    """Set cell width at XML level."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Set tcW
    tcW_nodes = tcPr.xpath('./w:tcW')
    if tcW_nodes:
        tcW = tcW_nodes[0]
        tcW.set(qn('w:w'), str(cm_to_twips(width_cm)))
        tcW.set(qn('w:type'), 'dxa')
    else:
        new_tcW = OxmlElement('w:tcW')
        new_tcW.set(qn('w:w'), str(cm_to_twips(width_cm)))
        new_tcW.set(qn('w:type'), 'dxa')
        tcPr.append(new_tcW)

    # Zero cell margins
    for node in tcPr.xpath('./w:tcMar'):
        tcPr.remove(node)
    tcMar_xml = parse_xml(
        r'<w:tcMar %s><w:left w:w="0" w:type="dxa"/><w:right w:w="0" w:type="dxa"/>'
        r'<w:top w:w="0" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/></w:tcMar>' % nsdecls("w")
    )
    tcPr.append(tcMar_xml)


def force_table_cell_widths(table, col_widths_cm):
    """Force cell widths for all cells."""
    for r in table.rows:
        for ci, cell in enumerate(r.cells):
            set_cell_tcPr_width(cell, col_widths_cm[ci])


def build_docx(labels, outpath, config):
    """Build Word document with labels."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(A4_W_CM)
    sec.page_height = Cm(A4_H_CM)
    sec.top_margin = Cm(config['margins_cm']['top'])
    sec.bottom_margin = Cm(config['margins_cm']['bottom'])
    sec.left_margin = Cm(config['margins_cm']['left'])
    sec.right_margin = Cm(config['margins_cm']['right'])

    pos = 0
    total = len(labels)

    while pos < total:
        # Calculate actual rows including gutters
        if config['has_row_gutters']:
            actual_rows = config['rows_per_page'] * 2 - 1
        else:
            actual_rows = config['rows_per_page']

        table = doc.add_table(rows=actual_rows, cols=len(config['col_widths_cm']))
        table.autofit = False
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Enforce XML-level layout & widths
        set_table_xml_enforced(table, config['col_widths_cm'])
        force_table_cell_widths(table, config['col_widths_cm'])

        # Set first-row widths
        for ci, w in enumerate(config['col_widths_cm']):
            table.rows[0].cells[ci].width = Cm(w)

        row_idx = 0
        for label_row in range(config['rows_per_page']):
            # Label row
            r = table.rows[row_idx]
            r.height = Cm(config['row_h_cm'])
            r.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            for ci, cell in enumerate(r.cells):
                if ci in config['label_col_idx'] and pos < total:
                    lab = labels[pos]

                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    try:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    except Exception:
                        pass

                    if not lab:
                        pos += 1
                        continue

                    p = cell.paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    # BHARAT header
                    run_header = p.add_run("BHARAT\n")
                    run_header.font.name = config['font']
                    run_header.font.size = Pt(config['font_size_header'])
                    run_header.bold = True

                    # Label code
                    run_text = p.add_run(lab)
                    run_text.font.name = config['font']
                    run_text.font.size = Pt(config['font_size'])

                    pos += 1

            row_idx += 1

            # Add gutter row if needed
            if config['has_row_gutters'] and label_row < config['rows_per_page'] - 1:
                gutter_row = table.rows[row_idx]
                gutter_row.height = Cm(config['row_gutter_h_cm'])
                gutter_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row_idx += 1

        if pos < total:
            doc.add_page_break()

    doc.save(str(outpath))
    return outpath


def generate_labels_for_codes(codes: List[str], output_dir: Path, date_str: str = "") -> List[Path]:
    """
    Main function to generate all label documents.
    
    Args:
        codes: List of participant codes
        output_dir: Directory to save documents
        date_str: Optional date string for filename suffix
    
    Returns:
        List of paths to generated .docx files
    """
    # Sort codes
    codes = sorted(codes)
    
    # Create label collections
    collections = create_label_collections(codes)
    
    # Create output directory
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Create filename suffix
    filename_suffix = f"_{date_str}" if date_str else ""
    
    generated_files = []
    
    # Generate cryovial labels (5 per row)
    cryo_filename = f'labels_cryovial{filename_suffix}.docx'
    cryo_path = build_docx(collections['cryovial'], output_dir / cryo_filename, CRYO_CONFIG)
    generated_files.append(cryo_path)
    
    # Generate normal labels (4 per row)
    for group_name in ['epigenetics', 'samples', 'edta', 'sst_fl_blood']:
        normal_filename = f'labels_{group_name}{filename_suffix}.docx'
        normal_path = build_docx(collections[group_name], output_dir / normal_filename, NORMAL_CONFIG)
        generated_files.append(normal_path)
    
    return generated_files
