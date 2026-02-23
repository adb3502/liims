"""
Label generator for BHARAT Study participants.

Generates A4-formatted Word documents with participant labels across 6 groups:
  - Cryovial (P1-P5): 5 per row, 17 rows/page, with row gutters
  - Urine (U1): 5 per row, same cryo layout
  - Epigenetics (E1-E4): 4 per row, 21 rows/page
  - Samples (CS1, R1, H1, +H2 for B-participants): 4 per row
  - EDTA (EDTA1-EDTA4): 4 per row
  - SST/Fl/Blood (SST1, SST2, Fl1, B1): 4 per row

Adapted from the original standalone label_generator.py script.
Uses LibreOffice headless for docx→pdf conversion.
"""

import io
import logging
import re
import subprocess
import tempfile
import zipfile
from pathlib import Path

logger = logging.getLogger(__name__)

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt

# 22 labels per participant, grouped by document type
LABEL_GROUPS = {
    "cryovial": ["P1", "P2", "P3", "P4", "P5"],
    "epigenetics": ["E1", "E2", "E3", "E4"],
    "samples": ["CS1", "R1", "H1", ""],  # blank slot becomes H2 for B-participants
    "edta": ["EDTA1", "EDTA2", "EDTA3", "EDTA4"],
    "sst_fl_blood": ["SST1", "SST2", "Fl1", "B1"],
    "urine": ["U1"],
}

# Cryovial layout: 5 labels per row with row gutters
CRYO_CONFIG = {
    "rows_per_page": 17,
    "row_h_cm": 1.3,
    "row_gutter_h_cm": 0.3,
    "col_widths_cm": [3.3, 0.3, 3.3, 0.3, 3.3, 0.3, 3.3, 0.3, 3.3],
    "label_col_idx": [0, 2, 4, 6, 8],
    "margins_cm": {"top": 0.5, "bottom": 0.5, "left": 0.5, "right": 0.5},
    "font": "Aptos",
    "font_size": 11,
    "font_size_header": 7,
    "has_row_gutters": True,
}

# Normal layout: 4 labels per row, no gutters
NORMAL_CONFIG = {
    "rows_per_page": 21,
    "row_h_cm": 1.25,
    "row_gutter_h_cm": 0,
    "col_widths_cm": [4.6, 0.5, 4.6, 0.5, 4.6, 0.5, 4.6],
    "label_col_idx": [0, 2, 4, 6],
    "margins_cm": {"top": 1.8, "bottom": 0.5, "left": 0.5, "right": 0.5},
    "font": "Aptos",
    "font_size": 11,
    "font_size_header": 7,
    "has_row_gutters": False,
}

A4_W_CM = 21.0
A4_H_CM = 29.7


def _cm_to_twips(cm: float) -> int:
    return int(cm * 567)


def _is_b_participant(participant: str) -> bool:
    """Check if participant code has 'B' (e.g., 1B-123, 2B-456)."""
    return bool(re.match(r"^\d+B-", participant))


def _create_label_collections(
    participants: list[str],
) -> dict[str, list[str]]:
    """Create collections of labels grouped by document type."""
    collections: dict[str, list[str]] = {k: [] for k in LABEL_GROUPS}

    for participant in participants:
        for group_name, suffixes in LABEL_GROUPS.items():
            for suffix in suffixes:
                if suffix:
                    collections[group_name].append(f"{participant}-{suffix}")
                elif _is_b_participant(participant):
                    collections[group_name].append(f"{participant}-H2")
                else:
                    collections[group_name].append("")

    return collections


def _set_table_xml_enforced(table, col_widths_cm: list[float]) -> None:
    """Add fixed layout + zero spacing + explicit widths at XML level."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Fixed layout
    try:
        layout_xml = parse_xml(
            r'<w:tblLayout %s w:type="fixed"/>' % nsdecls("w")
        )
        tblPr.append(layout_xml)
    except Exception:
        pass

    # Zero cell spacing
    try:
        spacing_xml = parse_xml(
            r'<w:tblCellSpacing %s w:w="0" w:type="dxa"/>' % nsdecls("w")
        )
        tblPr.append(spacing_xml)
    except Exception:
        pass

    # Table width
    total_twips = sum(_cm_to_twips(c) for c in col_widths_cm)
    try:
        tblW_xml = parse_xml(
            r'<w:tblW %s w:w="%d" w:type="dxa"/>' % (nsdecls("w"), total_twips)
        )
        tblPr.append(tblW_xml)
    except Exception:
        pass

    # Build tblGrid
    try:
        tblGrid = OxmlElement("w:tblGrid")
        for w in col_widths_cm:
            gridCol = OxmlElement("w:gridCol")
            gridCol.set(qn("w:w"), str(_cm_to_twips(w)))
            tblGrid.append(gridCol)
        tbl.append(tblGrid)
    except Exception:
        pass


def _set_cell_width(cell, width_cm: float) -> None:
    """Set cell width at XML level with zero margins."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcW_nodes = tcPr.xpath("./w:tcW")
    if tcW_nodes:
        tcW = tcW_nodes[0]
        tcW.set(qn("w:w"), str(_cm_to_twips(width_cm)))
        tcW.set(qn("w:type"), "dxa")
    else:
        new_tcW = OxmlElement("w:tcW")
        new_tcW.set(qn("w:w"), str(_cm_to_twips(width_cm)))
        new_tcW.set(qn("w:type"), "dxa")
        tcPr.append(new_tcW)

    # Zero cell margins
    for node in tcPr.xpath("./w:tcMar"):
        tcPr.remove(node)
    tcMar_xml = parse_xml(
        r"<w:tcMar %s>"
        r'<w:left w:w="0" w:type="dxa"/>'
        r'<w:right w:w="0" w:type="dxa"/>'
        r'<w:top w:w="0" w:type="dxa"/>'
        r'<w:bottom w:w="0" w:type="dxa"/>'
        r"</w:tcMar>" % nsdecls("w")
    )
    tcPr.append(tcMar_xml)


def _force_table_cell_widths(table, col_widths_cm: list[float]) -> None:
    for r in table.rows:
        for ci, cell in enumerate(r.cells):
            _set_cell_width(cell, col_widths_cm[ci])


def _build_docx(labels: list[str], config: dict) -> io.BytesIO:
    """Build a Word document with labels and return as BytesIO buffer."""
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(A4_W_CM)
    sec.page_height = Cm(A4_H_CM)
    sec.top_margin = Cm(config["margins_cm"]["top"])
    sec.bottom_margin = Cm(config["margins_cm"]["bottom"])
    sec.left_margin = Cm(config["margins_cm"]["left"])
    sec.right_margin = Cm(config["margins_cm"]["right"])

    pos = 0
    total = len(labels)

    while pos < total:
        if config["has_row_gutters"]:
            actual_rows = config["rows_per_page"] * 2 - 1
        else:
            actual_rows = config["rows_per_page"]

        table = doc.add_table(
            rows=actual_rows, cols=len(config["col_widths_cm"])
        )
        table.autofit = False
        table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        _set_table_xml_enforced(table, config["col_widths_cm"])
        _force_table_cell_widths(table, config["col_widths_cm"])

        for ci, w in enumerate(config["col_widths_cm"]):
            table.rows[0].cells[ci].width = Cm(w)

        row_idx = 0
        for label_row in range(config["rows_per_page"]):
            r = table.rows[row_idx]
            r.height = Cm(config["row_h_cm"])
            r.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

            for ci, cell in enumerate(r.cells):
                if ci in config["label_col_idx"] and pos < total:
                    lab = labels[pos]

                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    try:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    except Exception:
                        pass

                    if not lab:
                        pos += 1
                        continue

                    p = cell.paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    run_header = p.add_run("BHARAT\n")
                    run_header.font.name = config["font"]
                    run_header.font.size = Pt(config["font_size_header"])
                    run_header.bold = True

                    run_text = p.add_run(lab)
                    run_text.font.name = config["font"]
                    run_text.font.size = Pt(config["font_size"])

                    pos += 1

            row_idx += 1

            if (
                config["has_row_gutters"]
                and label_row < config["rows_per_page"] - 1
            ):
                gutter_row = table.rows[row_idx]
                gutter_row.height = Cm(config["row_gutter_h_cm"])
                gutter_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                row_idx += 1

        if pos < total:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


PDF_WORKER_DIR = Path("/data/pdf-worker")


def _convert_docx_files_to_pdf(
    docx_files: dict[str, bytes],
    timeout_seconds: int = 120,
) -> dict[str, bytes]:
    """
    Convert .docx files to PDF via the Windows worker (MS Word COM).

    Falls back to LibreOffice if the worker directory is not available
    (e.g. running outside Docker or worker not set up).

    Args:
        docx_files: dict of {filename: docx_bytes}

    Returns:
        dict of {filename_with_pdf_ext: pdf_bytes}
    """
    worker_dir = PDF_WORKER_DIR if PDF_WORKER_DIR.exists() else None

    if worker_dir:
        return _convert_via_word_worker(docx_files, worker_dir, timeout_seconds)
    else:
        logger.info("PDF worker dir not found, falling back to LibreOffice")
        return _convert_via_libreoffice(docx_files)


def _convert_via_word_worker(
    docx_files: dict[str, bytes],
    worker_dir: Path,
    timeout_seconds: int,
) -> dict[str, bytes]:
    """Send docx files to the Windows Word worker for conversion."""
    import uuid as _uuid
    import json
    import time as _time

    request_id = str(_uuid.uuid4())[:8]
    request_dir = worker_dir / request_id
    request_dir.mkdir(parents=True, exist_ok=True)

    # Write docx files
    filenames = []
    for name, data in docx_files.items():
        (request_dir / name).write_bytes(data)
        filenames.append(name)

    # Write request file (signals the worker)
    request_file = worker_dir / f"{request_id}.request.json"
    request_file.write_text(json.dumps({
        "request_id": request_id,
        "files": filenames,
    }))

    logger.info("PDF request %s: %d files, waiting for worker...", request_id, len(filenames))

    # Wait for worker to finish
    done_file = request_dir / "done.json"
    for _ in range(timeout_seconds * 2):  # Check every 0.5s
        if done_file.exists():
            break
        _time.sleep(0.5)
    else:
        # Timeout — clean up and fall back to LibreOffice
        logger.warning("PDF worker timeout for request %s, falling back to LibreOffice", request_id)
        _cleanup_request_dir(request_dir, request_file)
        return _convert_via_libreoffice(docx_files)

    # Read PDFs
    results: dict[str, bytes] = {}
    for name in filenames:
        pdf_name = name.replace(".docx", ".pdf")
        pdf_path = request_dir / pdf_name
        if pdf_path.exists():
            results[pdf_name] = pdf_path.read_bytes()
        else:
            logger.warning("PDF not found for %s", name)

    # Clean up
    _cleanup_request_dir(request_dir)
    return results


def _cleanup_request_dir(request_dir: Path, request_file: Path | None = None) -> None:
    """Remove request directory and request file."""
    try:
        if request_file and request_file.exists():
            request_file.unlink()
        if request_dir.exists():
            for f in request_dir.iterdir():
                f.unlink()
            request_dir.rmdir()
    except Exception as e:
        logger.warning("Cleanup error: %s", e)


def _convert_via_libreoffice(docx_files: dict[str, bytes]) -> dict[str, bytes]:
    """Fallback: convert docx to PDF using LibreOffice headless."""
    results: dict[str, bytes] = {}

    for name, data in docx_files.items():
        with tempfile.TemporaryDirectory() as tmpdir:
            tmpdir_path = Path(tmpdir)
            docx_path = tmpdir_path / "input.docx"
            docx_path.write_bytes(data)

            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(tmpdir_path),
                    str(docx_path),
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )

            pdf_path = tmpdir_path / "input.pdf"
            if result.returncode == 0 and pdf_path.exists():
                pdf_name = name.replace(".docx", ".pdf")
                results[pdf_name] = pdf_path.read_bytes()
            else:
                logger.error("LibreOffice failed for %s: %s", name, result.stderr)

    return results


def generate_label_zip(
    participant_codes: list[str],
    date_str: str = "",
    output_format: str = "docx",
) -> io.BytesIO:
    """
    Generate all 6 label documents and return as a ZIP file in memory.

    Args:
        participant_codes: List of participant codes (e.g. ["1A-001", "2B-045"])
        date_str: Optional date string appended to filenames
        output_format: "docx" or "pdf"

    Returns:
        BytesIO containing a ZIP with 6 files
    """
    codes = sorted(participant_codes)
    collections = _create_label_collections(codes)
    suffix = f"_{date_str}" if date_str else ""
    ext = output_format if output_format in ("docx", "pdf") else "docx"

    # Build all docx buffers first
    group_order = [
        ("cryovial", CRYO_CONFIG),
        ("urine", CRYO_CONFIG),
        ("epigenetics", NORMAL_CONFIG),
        ("samples", NORMAL_CONFIG),
        ("edta", NORMAL_CONFIG),
        ("sst_fl_blood", NORMAL_CONFIG),
    ]

    # Build all docx buffers
    docx_outputs: dict[str, bytes] = {}
    for group_name, config in group_order:
        docx_buf = _build_docx(collections[group_name], config)
        docx_name = f"labels_{group_name}{suffix}.docx"
        docx_outputs[docx_name] = docx_buf.read()

    # Convert to PDF if requested
    if ext == "pdf":
        pdf_outputs = _convert_docx_files_to_pdf(docx_outputs)
        files_to_zip = pdf_outputs
    else:
        files_to_zip = docx_outputs

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename, data in files_to_zip.items():
            zf.writestr(filename, data)

    zip_buf.seek(0)
    return zip_buf


def generate_single_label_doc(
    participant_codes: list[str],
    group: str,
) -> io.BytesIO:
    """
    Generate a single label group document.

    Args:
        participant_codes: List of participant codes
        group: One of 'cryovial', 'epigenetics', 'samples', 'edta', 'sst_fl_blood'

    Returns:
        BytesIO containing a .docx file
    """
    if group not in LABEL_GROUPS:
        raise ValueError(f"Invalid group '{group}'. Must be one of: {list(LABEL_GROUPS.keys())}")

    codes = sorted(participant_codes)
    collections = _create_label_collections(codes)
    config = CRYO_CONFIG if group in ("cryovial", "urine") else NORMAL_CONFIG
    return _build_docx(collections[group], config)
